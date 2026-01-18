"""
Resume Exporter - Generates PDF and DOCX files from resume data
"""
import logging
from typing import Dict, Any
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

logger = logging.getLogger(__name__)


def export_to_pdf(resume_data: Dict[str, Any]) -> bytes:
    """
    Generate PDF from resume data
    Returns PDF bytes
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            "ResumeTitle",
            parent=styles["Heading1"],
            fontSize=24,
            textColor=colors.HexColor("#1a2234"),
            spaceAfter=12,
            alignment=1,  # Center
        )
        
        # Heading style
        heading_style = ParagraphStyle(
            "ResumeHeading",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=colors.HexColor("#366092"),
            spaceAfter=6,
            spaceBefore=12,
        )
        
        # Body style
        body_style = styles["Normal"]
        body_style.fontSize = 10
        body_style.leading = 12
        
        # Name and contact
        name = resume_data.get("title", "Resume")
        story.append(Paragraph(name, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Summary
        if resume_data.get("summary"):
            story.append(Paragraph("Summary", heading_style))
            story.append(Paragraph(resume_data["summary"], body_style))
            story.append(Spacer(1, 0.15*inch))
        
        # Experience
        if resume_data.get("experience"):
            story.append(Paragraph("Experience", heading_style))
            for exp in resume_data["experience"]:
                company = exp.get("company", "")
                role = exp.get("role", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "")
                description = exp.get("description", "")
                
                exp_text = f"<b>{role}</b> at <b>{company}</b>"
                if start_date or end_date:
                    exp_text += f" ({start_date} - {end_date})"
                story.append(Paragraph(exp_text, body_style))
                if description:
                    story.append(Paragraph(description, body_style))
                story.append(Spacer(1, 0.1*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Education
        if resume_data.get("education"):
            story.append(Paragraph("Education", heading_style))
            for edu in resume_data["education"]:
                institution = edu.get("institution", "")
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                graduation_date = edu.get("graduation_date", "")
                
                edu_text = f"<b>{degree}</b>"
                if field:
                    edu_text += f" in {field}"
                if institution:
                    edu_text += f" - {institution}"
                if graduation_date:
                    edu_text += f" ({graduation_date})"
                story.append(Paragraph(edu_text, body_style))
                story.append(Spacer(1, 0.1*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Skills
        if resume_data.get("skills"):
            story.append(Paragraph("Skills", heading_style))
            skills_text = ", ".join(resume_data["skills"])
            story.append(Paragraph(skills_text, body_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Projects
        if resume_data.get("projects"):
            story.append(Paragraph("Projects", heading_style))
            for project in resume_data["projects"]:
                name = project.get("name", "")
                description = project.get("description", "")
                story.append(Paragraph(f"<b>{name}</b>", body_style))
                if description:
                    story.append(Paragraph(description, body_style))
                story.append(Spacer(1, 0.1*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Certifications
        if resume_data.get("certifications"):
            story.append(Paragraph("Certifications", heading_style))
            for cert in resume_data["certifications"]:
                name = cert.get("name", "")
                issuer = cert.get("issuer", "")
                date = cert.get("date", "")
                cert_text = f"<b>{name}</b>"
                if issuer:
                    cert_text += f" - {issuer}"
                if date:
                    cert_text += f" ({date})"
                story.append(Paragraph(cert_text, body_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        pdf_bytes = buffer.read()
        logger.info(f"PDF exported successfully, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"Error exporting PDF: {e}", exc_info=True)
        raise ValueError(f"Failed to export PDF: {str(e)}")


def export_to_docx(resume_data: Dict[str, Any]) -> bytes:
    """
    Generate DOCX from resume data
    Returns DOCX bytes
    """
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(resume_data.get("title", "Resume"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if resume_data.get("summary"):
            doc.add_heading("Summary", 1)
            doc.add_paragraph(resume_data["summary"])
        
        # Experience
        if resume_data.get("experience"):
            doc.add_heading("Experience", 1)
            for exp in resume_data["experience"]:
                company = exp.get("company", "")
                role = exp.get("role", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "")
                description = exp.get("description", "")
                
                p = doc.add_paragraph()
                p.add_run(f"{role} at {company}").bold = True
                if start_date or end_date:
                    p.add_run(f" ({start_date} - {end_date})")
                if description:
                    doc.add_paragraph(description, style='List Bullet')
        
        # Education
        if resume_data.get("education"):
            doc.add_heading("Education", 1)
            for edu in resume_data["education"]:
                institution = edu.get("institution", "")
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                graduation_date = edu.get("graduation_date", "")
                
                p = doc.add_paragraph()
                p.add_run(degree).bold = True
                if field:
                    p.add_run(f" in {field}")
                if institution:
                    p.add_run(f" - {institution}")
                if graduation_date:
                    p.add_run(f" ({graduation_date})")
        
        # Skills
        if resume_data.get("skills"):
            doc.add_heading("Skills", 1)
            skills_text = ", ".join(resume_data["skills"])
            doc.add_paragraph(skills_text)
        
        # Projects
        if resume_data.get("projects"):
            doc.add_heading("Projects", 1)
            for project in resume_data["projects"]:
                name = project.get("name", "")
                description = project.get("description", "")
                p = doc.add_paragraph()
                p.add_run(name).bold = True
                if description:
                    doc.add_paragraph(description, style='List Bullet')
        
        # Certifications
        if resume_data.get("certifications"):
            doc.add_heading("Certifications", 1)
            for cert in resume_data["certifications"]:
                name = cert.get("name", "")
                issuer = cert.get("issuer", "")
                date = cert.get("date", "")
                p = doc.add_paragraph()
                p.add_run(name).bold = True
                if issuer:
                    p.add_run(f" - {issuer}")
                if date:
                    p.add_run(f" ({date})")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
        logger.info(f"DOCX exported successfully, size: {len(docx_bytes)} bytes")
        return docx_bytes
        
    except Exception as e:
        logger.error(f"Error exporting DOCX: {e}", exc_info=True)
        raise ValueError(f"Failed to export DOCX: {str(e)}")
