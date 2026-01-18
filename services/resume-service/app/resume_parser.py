"""
Resume Parser - Extracts structured data from PDF/DOCX files
"""
import logging
from typing import Dict, List, Any
import PyPDF2
from docx import Document
import io

logger = logging.getLogger(__name__)


def parse_pdf(file_content: bytes) -> Dict[str, Any]:
    """
    Parse PDF resume and extract text
    Returns structured data dictionary
    """
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text() + "\n"
        
        # Basic parsing - extract structured data from text
        parsed = _extract_structured_data(text_content)
        logger.info(f"PDF parsed successfully, extracted {len(parsed.get('experience', []))} experiences")
        return parsed
        
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}", exc_info=True)
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_content: bytes) -> Dict[str, Any]:
    """
    Parse DOCX resume and extract text
    Returns structured data dictionary
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        
        # Basic parsing - extract structured data from text
        parsed = _extract_structured_data(text_content)
        logger.info(f"DOCX parsed successfully, extracted {len(parsed.get('experience', []))} experiences")
        return parsed
        
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}", exc_info=True)
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _extract_structured_data(text: str) -> Dict[str, Any]:
    """
    Extract structured data from raw text
    This is a basic implementation - can be enhanced with NLP/ML
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    parsed = {
        "name": "",
        "email": "",
        "phone": "",
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }
    
    # Extract email
    import re
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        parsed["email"] = emails[0]
    
    # Extract phone
    phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
    phones = re.findall(phone_pattern, text)
    if phones:
        parsed["phone"] = ''.join(phones[0])
    
    # Try to identify sections
    current_section = None
    experience_keywords = ["experience", "work", "employment", "career", "professional"]
    education_keywords = ["education", "academic", "university", "college", "degree"]
    skills_keywords = ["skills", "technical", "competencies", "expertise"]
    project_keywords = ["projects", "portfolio", "work samples"]
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Detect section headers
        if any(keyword in line_lower for keyword in experience_keywords) and len(line) < 50:
            current_section = "experience"
            continue
        elif any(keyword in line_lower for keyword in education_keywords) and len(line) < 50:
            current_section = "education"
            continue
        elif any(keyword in line_lower for keyword in skills_keywords) and len(line) < 50:
            current_section = "skills"
            continue
        elif any(keyword in line_lower for keyword in project_keywords) and len(line) < 50:
            current_section = "projects"
            continue
        
        # Extract name (usually first line or after header)
        if i < 3 and not parsed["name"] and len(line) < 100 and not any(char.isdigit() for char in line):
            parsed["name"] = line
        
        # Extract summary (usually after name, before sections)
        if not current_section and i < 10 and len(line) > 50:
            if not parsed["summary"]:
                parsed["summary"] = line
            else:
                parsed["summary"] += " " + line
        
        # Extract experience entries (basic pattern matching)
        if current_section == "experience":
            # Look for date patterns or company names
            if re.search(r'\d{4}|\d{1,2}/\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', line):
                # Potential experience entry
                if len(parsed["experience"]) == 0 or len(parsed["experience"][-1].get("description", "")) > 50:
                    parsed["experience"].append({
                        "company": line,
                        "role": "",
                        "start_date": "",
                        "end_date": "",
                        "description": "",
                    })
        
        # Extract skills (comma-separated or bullet points)
        if current_section == "skills":
            # Split by comma, semicolon, or bullet
            skill_items = re.split(r'[,;•·]', line)
            for skill in skill_items:
                skill = skill.strip()
                if skill and len(skill) < 50:
                    parsed["skills"].append(skill)
    
    # Clean up skills (remove duplicates, limit length)
    parsed["skills"] = list(dict.fromkeys(parsed["skills"]))[:50]  # Max 50 skills
    
    return parsed
